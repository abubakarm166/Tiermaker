"""
Convert docs/API_REFERENCE.md to docs/API_REFERENCE.docx for Google Docs import.
Run: python docs/build_api_reference_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def clean_heading(s: str) -> str:
    return re.sub(r"\*\*", "", s).strip()


def is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def is_table_sep(line: str) -> bool:
    stripped = line.strip().strip("|")
    if not stripped:
        return False
    return all(c in "-: |" for c in stripped)


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def simplify_cell(text: str) -> str:
    return re.sub(r"\*\*([^*]+)\*\*", r"\1", text)


def add_mixed_paragraph(doc: Document, text: str) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
        else:
            p.add_run(part)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = simplify_cell(row[j]) if j < len(row) else ""
            table.rows[i].cells[j].text = cell_text


def main() -> None:
    md_path = Path(__file__).resolve().parent / "API_REFERENCE.md"
    out_path = md_path.with_suffix(".docx")
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_lines))
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if is_table_row(line):
            rows: list[list[str]] = []
            while i < len(lines) and is_table_row(lines[i]):
                if not is_table_sep(lines[i]):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if line.startswith("# ") and not line.startswith("##"):
            doc.add_heading(clean_heading(line[2:]), level=0)
        elif line.startswith("## "):
            doc.add_heading(clean_heading(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(clean_heading(line[4:]), level=2)
        elif not line.strip():
            pass
        else:
            add_mixed_paragraph(doc, line)

        i += 1

    try:
        doc.save(out_path)
        print(f"Wrote {out_path}")
    except PermissionError:
        fallback = out_path.with_name(f"{out_path.stem}_generated.docx")
        doc.save(fallback)
        print(
            f"Could not overwrite {out_path} (file may be open). "
            f"Wrote {fallback} instead."
        )


if __name__ == "__main__":
    main()
